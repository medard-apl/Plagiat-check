"""
Génère deux fichiers .docx factices pour tester le pipeline complet :

- doc_a.docx : le document "source" (comme s'il existait déjà quelque part).
- doc_b.docx : un document qui reprend le contenu de A avec 3 cas différents :
    - un paragraphe COPIÉ mot pour mot (doit être détecté par le fingerprinting)
    - un paragraphe PARAPHRASÉ (doit être détecté par les embeddings, pas par le fingerprinting)
    - un paragraphe ORIGINAL, sans rapport (ne doit rien détecter du tout)

On génère de vrais .docx (pas juste des chaînes Python) pour tester le pipeline
en conditions réelles, extraction comprise, et pas seulement la logique de
comparaison.
"""

from pathlib import Path
import docx

FIXTURES_DIR = Path(__file__).parent

PARAGRAPH_COPIED = (
    "Les algorithmes d'apprentissage automatique permettent aujourd'hui de résoudre "
    "des problèmes complexes qui semblaient inaccessibles il y a encore une décennie."
)

PARAGRAPH_ORIGINAL_TO_PARAPHRASE = (
    "La détection de plagiat repose sur la comparaison de segments de texte à l'aide "
    "de techniques complémentaires, combinant analyse lexicale exacte et analyse sémantique."
)

PARAGRAPH_PARAPHRASED = (
    "Pour repérer un contenu copié, on combine plusieurs méthodes de comparaison de texte, "
    "certaines fondées sur la correspondance de mots précis, d'autres sur le sens général."
)

PARAGRAPH_UNRELATED_A = (
    "Le climat de la région étudiée se caractérise par des précipitations abondantes "
    "durant la saison des pluies, entre juin et septembre."
)

PARAGRAPH_UNRELATED_B = (
    "Les félins domestiques passent en moyenne seize heures par jour à dormir, "
    "un comportement hérité de leurs ancêtres chasseurs nocturnes."
)


def build_doc_a() -> None:
    document = docx.Document()
    document.add_paragraph(PARAGRAPH_COPIED)
    document.add_paragraph(PARAGRAPH_ORIGINAL_TO_PARAPHRASE)
    document.add_paragraph(PARAGRAPH_UNRELATED_A)
    document.add_paragraph("Bibliographie")
    document.add_paragraph("Dupont, J. (2020). Introduction au machine learning. Éditions Exemple.")
    document.save(FIXTURES_DIR / "doc_a.docx")


def build_doc_b() -> None:
    document = docx.Document()
    document.add_paragraph(PARAGRAPH_COPIED)  # copie exacte du paragraphe de A
    document.add_paragraph(PARAGRAPH_PARAPHRASED)  # reformulation du 2e paragraphe de A
    document.add_paragraph(PARAGRAPH_UNRELATED_B)  # sans rapport avec A
    document.add_paragraph("Bibliographie")
    document.add_paragraph("Martin, P. (2021). Les félins domestiques. Éditions Exemple.")
    document.save(FIXTURES_DIR / "doc_b.docx")


if __name__ == "__main__":
    build_doc_a()
    build_doc_b()
    print(f"Fichiers générés dans {FIXTURES_DIR} : doc_a.docx, doc_b.docx")